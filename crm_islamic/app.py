import os
import sys
import csv
import io
import uuid
import copy
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document as DocxDocument
from datetime import date, datetime, timedelta
from dateutil.relativedelta import relativedelta
from werkzeug.utils import secure_filename

from flask import (
    Flask, render_template, request, redirect, url_for,
    flash, jsonify, send_file, abort, Response
)
from flask_wtf.csrf import CSRFProtect, generate_csrf
from sqlalchemy import func, desc, asc, extract

from .models import db, Client, Contract, PaymentSchedule, Payment, Guarantor, Backup, Document
from .forms import (
    ClientForm, ContractForm, PaymentForm, GuarantorForm,
    EMPLOYMENT_CHOICES, MARITAL_CHOICES, EDUCATION_CHOICES,
    CREDIT_HISTORY_CHOICES, STATUS_CHOICES, ITEM_CATEGORY_CHOICES,
    PAYMENT_METHOD_CHOICES
)
from . import backup as backup_module

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'webp', 'pdf', 'doc', 'docx', 'xls', 'xlsx', 'heic'}


def _allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def create_app(config=None):
    # Resolve template/static paths when running as PyInstaller bundle
    if getattr(sys, 'frozen', False):
        base_dir = sys._MEIPASS
    else:
        base_dir = os.path.dirname(os.path.abspath(__file__))

    app = Flask(
        __name__,
        template_folder=os.path.join(base_dir, 'templates'),
        static_folder=os.path.join(base_dir, 'static') if os.path.isdir(os.path.join(base_dir, 'static')) else None,
    )
    app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'islamic-crm-secret-key-2024-murabaha')

    # DB lives next to the executable when frozen, otherwise next to this file
    if getattr(sys, 'frozen', False):
        db_dir = os.environ.get('CRM_DB_DIR', os.path.dirname(sys.executable))
    else:
        db_dir = os.path.dirname(os.path.abspath(__file__))
    db_path = os.path.join(db_dir, 'crm_islamic.db')
    app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{db_path}'
    app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
    app.config['WTF_CSRF_ENABLED'] = True

    app.config['DB_DIR'] = db_dir
    upload_dir = os.path.join(db_dir, 'uploads')
    os.makedirs(upload_dir, exist_ok=True)
    app.config['UPLOAD_FOLDER'] = upload_dir
    app.config['MAX_CONTENT_LENGTH'] = 32 * 1024 * 1024  # 32 MB

    if config:
        app.config.update(config)

    db.init_app(app)
    csrf = CSRFProtect(app)

    # ── helpers ────────────────────────────────────────────────────────────────

    def _get_upload_dir():
        return app.config['UPLOAD_FOLDER']

    def _save_inline_guarantors(client_id, form_data, clear_existing=False):
        if clear_existing:
            db.session.execute(
                db.delete(Guarantor).where(
                    Guarantor.client_id == client_id,
                    Guarantor.contract_id.is_(None)
                )
            )
        count = form_data.get('guarantor_count', 0, type=int)
        for i in range(count):
            first_name = form_data.get(f'g_{i}_first_name', '').strip()
            last_name = form_data.get(f'g_{i}_last_name', '').strip()
            if not first_name or not last_name:
                continue
            g = Guarantor(
                client_id=client_id,
                contract_id=None,
                first_name=first_name,
                last_name=last_name,
                middle_name=form_data.get(f'g_{i}_middle_name', '').strip() or None,
                phone=form_data.get(f'g_{i}_phone', '').strip() or None,
                phone2=form_data.get(f'g_{i}_phone2', '').strip() or None,
                email=form_data.get(f'g_{i}_email', '').strip() or None,
                relationship=form_data.get(f'g_{i}_relationship', '') or None,
                guarantor_type=form_data.get(f'g_{i}_guarantor_type', 'personal'),
                passport_series=form_data.get(f'g_{i}_passport_series', '').strip() or None,
                passport_number=form_data.get(f'g_{i}_passport_number', '').strip() or None,
                passport_issued_by=form_data.get(f'g_{i}_passport_issued_by', '').strip() or None,
                inn=form_data.get(f'g_{i}_inn', '').strip() or None,
                snils=form_data.get(f'g_{i}_snils', '').strip() or None,
                address_registration=form_data.get(f'g_{i}_address_registration', '').strip() or None,
                address_actual=form_data.get(f'g_{i}_address_actual', '').strip() or None,
                employer_name=form_data.get(f'g_{i}_employer_name', '').strip() or None,
                employer_phone=form_data.get(f'g_{i}_employer_phone', '').strip() or None,
                position=form_data.get(f'g_{i}_position', '').strip() or None,
                employment_type=form_data.get(f'g_{i}_employment_type', '') or None,
                monthly_income=float(form_data.get(f'g_{i}_monthly_income') or 0),
                property_description=form_data.get(f'g_{i}_property_description', '').strip() or None,
                notes=form_data.get(f'g_{i}_notes', '').strip() or None,
            )
            for attr, field in (('passport_issued_date', f'g_{i}_passport_issued_date'),
                                 ('birth_date', f'g_{i}_birth_date')):
                val = form_data.get(field, '').strip()
                if val:
                    try:
                        setattr(g, attr, date.fromisoformat(val))
                    except Exception:
                        pass
            db.session.add(g)

    def generate_contract_number():
        year = datetime.utcnow().year
        last = db.session.execute(
            db.select(Contract).where(
                Contract.contract_number.like(f'МУР-{year}-%')
            ).order_by(desc(Contract.id))
        ).scalars().first()
        if last and last.contract_number:
            try:
                seq = int(last.contract_number.split('-')[-1]) + 1
            except Exception:
                seq = 1
        else:
            seq = 1
        return f'МУР-{year}-{seq:03d}'

    def generate_schedule(contract):
        """Delete existing schedule rows and create fresh ones."""
        db.session.execute(
            db.delete(PaymentSchedule).where(PaymentSchedule.contract_id == contract.id)
        )
        db.session.flush()

        if not contract.first_payment_date or not contract.months:
            return

        base = contract.first_payment_date
        for i in range(contract.months):
            due = base + relativedelta(months=i)
            ps = PaymentSchedule(
                contract_id=contract.id,
                installment_num=i + 1,
                due_date=due,
                amount=round(contract.monthly_payment, 2),
                status='pending',
                paid_amount=0.0,
            )
            db.session.add(ps)

        contract.last_payment_date = base + relativedelta(months=contract.months - 1)
        db.session.flush()

    def update_overdue_statuses():
        today = date.today()
        # Mark overdue schedule items
        db.session.execute(
            db.update(PaymentSchedule).where(
                PaymentSchedule.due_date < today,
                PaymentSchedule.status.in_(['pending', 'partial'])
            ).values(status='overdue')
        )
        # Mark contracts as overdue if they have overdue schedules
        overdue_contract_ids = db.session.execute(
            db.select(PaymentSchedule.contract_id).where(
                PaymentSchedule.status == 'overdue'
            ).distinct()
        ).scalars().all()

        if overdue_contract_ids:
            db.session.execute(
                db.update(Contract).where(
                    Contract.id.in_(overdue_contract_ids),
                    Contract.status == 'active'
                ).values(status='overdue')
            )
        db.session.commit()

    @app.before_request
    def before_each_request():
        update_overdue_statuses()

    # ── context processors ────────────────────────────────────────────────────

    @app.context_processor
    def inject_globals():
        today = date.today()
        overdue_count = db.session.execute(
            db.select(func.count(Contract.id)).where(Contract.status == 'overdue')
        ).scalar() or 0
        return dict(today=today, nav_overdue_count=overdue_count, csrf_token=generate_csrf)

    # ── error handlers ────────────────────────────────────────────────────────

    @app.errorhandler(404)
    def not_found(e):
        return render_template('errors/404.html'), 404

    @app.errorhandler(500)
    def server_error(e):
        return render_template('errors/500.html'), 500

    # ══════════════════════════════════════════════════════════════════════════
    # DASHBOARD
    # ══════════════════════════════════════════════════════════════════════════

    @app.route('/')
    def dashboard():
        today = date.today()
        total_clients = db.session.execute(
            db.select(func.count(Client.id)).where(Client.status != 'blacklisted')
        ).scalar() or 0

        active_contracts = db.session.execute(
            db.select(func.count(Contract.id)).where(Contract.status.in_(['active', 'overdue']))
        ).scalar() or 0

        portfolio = db.session.execute(
            db.select(func.coalesce(func.sum(Contract.financed_amount), 0)).where(
                Contract.status.in_(['active', 'overdue'])
            )
        ).scalar() or 0

        overdue_contracts = db.session.execute(
            db.select(func.count(Contract.id)).where(Contract.status == 'overdue')
        ).scalar() or 0

        # Today's collections
        today_collected = db.session.execute(
            db.select(func.coalesce(func.sum(Payment.amount), 0)).where(
                Payment.payment_date == today
            )
        ).scalar() or 0

        # Upcoming payments next 7 days
        upcoming = db.session.execute(
            db.select(PaymentSchedule).where(
                PaymentSchedule.due_date >= today,
                PaymentSchedule.due_date <= today + timedelta(days=7),
                PaymentSchedule.status.in_(['pending', 'partial'])
            ).order_by(asc(PaymentSchedule.due_date)).limit(20)
        ).scalars().all()

        # Overdue contracts with days overdue
        overdue_list = db.session.execute(
            db.select(Contract).where(Contract.status == 'overdue').limit(10)
        ).scalars().all()

        return render_template('dashboard.html',
            total_clients=total_clients,
            active_contracts=active_contracts,
            portfolio=portfolio,
            overdue_contracts=overdue_contracts,
            today_collected=today_collected,
            upcoming=upcoming,
            overdue_list=overdue_list,
            today=today,
        )

    # ══════════════════════════════════════════════════════════════════════════
    # CLIENTS
    # ══════════════════════════════════════════════════════════════════════════

    @app.route('/clients')
    def clients_list():
        clients = db.session.execute(
            db.select(Client).order_by(desc(Client.created_at))
        ).scalars().all()

        # Export CSV
        if request.args.get('export') == 'csv':
            def generate():
                output = io.StringIO()
                w = csv.writer(output)
                w.writerow(['ID', 'Фамилия', 'Имя', 'Отчество', 'Телефон', 'Статус', 'Кредитный балл', 'Договоры'])
                for c in clients:
                    w.writerow([c.id, c.last_name, c.first_name, c.middle_name or '',
                                 c.phone or '', c.status, c.credit_score,
                                 c.contracts.count()])
                    output.seek(0)
                    yield output.read()
                    output.seek(0)
                    output.truncate()
            return Response(generate(), mimetype='text/csv',
                            headers={'Content-Disposition': 'attachment;filename=clients.csv'})

        return render_template('clients/list.html', clients=clients)

    @app.route('/clients/new', methods=['GET', 'POST'])
    def client_new():
        form = ClientForm()
        if form.validate_on_submit():
            client = Client()
            _populate_client(client, form)
            db.session.add(client)
            db.session.flush()
            _save_inline_guarantors(client.id, request.form)
            db.session.commit()
            flash(f'Клиент {client.full_name} добавлен', 'success')
            return redirect(url_for('client_detail', client_id=client.id))
        return render_template('clients/form.html', form=form, title='Новый клиент', client=None,
                               client_guarantors=[])

    @app.route('/clients/<int:client_id>')
    def client_detail(client_id):
        client = db.session.get(Client, client_id) or abort(404)
        contracts = db.session.execute(
            db.select(Contract).where(Contract.client_id == client_id).order_by(desc(Contract.created_at))
        ).scalars().all()
        guarantor_records = db.session.execute(
            db.select(Guarantor).where(Guarantor.client_id == client_id)
        ).scalars().all()
        documents = db.session.execute(
            db.select(Document).where(
                Document.entity_type == 'client', Document.entity_id == client_id
            ).order_by(desc(Document.uploaded_at))
        ).scalars().all()

        # Полный график платежей по всем договорам клиента
        contract_ids = [c.id for c in contracts]
        all_schedule = []
        if contract_ids:
            schedule_rows = db.session.execute(
                db.select(PaymentSchedule)
                .where(PaymentSchedule.contract_id.in_(contract_ids))
                .order_by(asc(PaymentSchedule.due_date))
            ).scalars().all()
            # Индекс договоров для быстрого доступа
            contracts_map = {c.id: c for c in contracts}
            for s in schedule_rows:
                c = contracts_map.get(s.contract_id)
                all_schedule.append({
                    'id': s.id,
                    'contract_id': s.contract_id,
                    'contract_number': c.contract_number if c else '—',
                    'item_name': c.item_name if c else '—',
                    'installment_num': s.installment_num,
                    'due_date': s.due_date.strftime('%d.%m.%Y') if s.due_date else '—',
                    'due_date_iso': s.due_date.isoformat() if s.due_date else '',
                    'amount': float(s.amount),
                    'paid_amount': float(s.paid_amount or 0),
                    'remaining': float(max(0, s.amount - (s.paid_amount or 0))),
                    'status': s.status,
                })

        return render_template('clients/detail.html', client=client,
                               contracts=contracts, guarantor_records=guarantor_records,
                               documents=documents, all_schedule=all_schedule)

    @app.route('/clients/<int:client_id>/edit', methods=['GET', 'POST'])
    def client_edit(client_id):
        client = db.session.get(Client, client_id) or abort(404)
        form = ClientForm(obj=client)
        if form.validate_on_submit():
            _populate_client(client, form)
            client.updated_at = datetime.utcnow()
            _save_inline_guarantors(client.id, request.form, clear_existing=True)
            db.session.commit()
            flash('Данные клиента обновлены', 'success')
            return redirect(url_for('client_detail', client_id=client.id))
        existing_guarantors = db.session.execute(
            db.select(Guarantor).where(
                Guarantor.client_id == client_id, Guarantor.contract_id.is_(None)
            )
        ).scalars().all()
        return render_template('clients/form.html', form=form, title='Редактировать клиента',
                               client=client, client_guarantors=existing_guarantors)

    @app.route('/clients/<int:client_id>/delete', methods=['POST'])
    def client_delete(client_id):
        client = db.session.get(Client, client_id) or abort(404)
        client.status = 'inactive'
        db.session.commit()
        flash(f'Клиент {client.full_name} деактивирован', 'warning')
        return redirect(url_for('clients_list'))

    def _populate_client(client, form):
        for field in form:
            if field.name not in ('submit', 'csrf_token'):
                try:
                    setattr(client, field.name, field.data)
                except Exception:
                    pass

    # ══════════════════════════════════════════════════════════════════════════
    # GUARANTORS
    # ══════════════════════════════════════════════════════════════════════════

    @app.route('/guarantors/new', methods=['GET', 'POST'])
    def guarantor_new():
        contract_id = request.args.get('contract_id', type=int) or request.form.get('contract_id', type=int)
        client_id = request.args.get('client_id', type=int) or request.form.get('client_id', type=int)
        contract = db.session.get(Contract, contract_id) if contract_id else None
        client_obj = db.session.get(Client, client_id) if (client_id and not contract) else None
        if not contract and not client_obj:
            abort(400)
        form = GuarantorForm()
        if request.method == 'GET':
            form.contract_id.data = str(contract_id) if contract_id else ''
            form.client_id.data = str(client_id) if client_id else (str(contract.client_id) if contract else '')
        if form.validate_on_submit():
            g = Guarantor(
                contract_id=contract_id,
                client_id=int(form.client_id.data) if form.client_id.data else (contract.client_id if contract else None),
                relationship=form.relationship.data,
                last_name=form.last_name.data,
                first_name=form.first_name.data,
                middle_name=form.middle_name.data or None,
                birth_date=form.birth_date.data,
                gender=form.gender.data or None,
                phone=form.phone.data or None,
                phone2=form.phone2.data or None,
                email=form.email.data or None,
                passport_series=form.passport_series.data or None,
                passport_number=form.passport_number.data or None,
                passport_issued_by=form.passport_issued_by.data or None,
                passport_issued_date=form.passport_issued_date.data,
                inn=form.inn.data or None,
                snils=form.snils.data or None,
                address_registration=form.address_registration.data or None,
                address_actual=form.address_actual.data or None,
                employer_name=form.employer_name.data or None,
                employer_phone=form.employer_phone.data or None,
                position=form.position.data or None,
                employment_type=form.employment_type.data or None,
                work_experience_months=form.work_experience_months.data,
                monthly_income=form.monthly_income.data or 0,
                guarantor_type=form.guarantor_type.data,
                property_description=form.property_description.data or None,
                notes=form.notes.data or None,
            )
            db.session.add(g)
            db.session.commit()
            flash('Поручитель добавлен', 'success')
            if contract:
                return redirect(url_for('contract_detail', contract_id=contract_id))
            return redirect(url_for('client_detail', client_id=client_id))
        return render_template('guarantors/form.html', form=form, contract=contract,
                               client_obj=client_obj, title='Новый поручитель')

    @app.route('/guarantors/<int:g_id>/edit', methods=['GET', 'POST'])
    def guarantor_edit(g_id):
        g = db.session.get(Guarantor, g_id) or abort(404)
        contract = db.session.get(Contract, g.contract_id) if g.contract_id else None
        client_obj = db.session.get(Client, g.client_id) if (g.client_id and not contract) else None
        form = GuarantorForm(obj=g)
        if form.validate_on_submit():
            for field in form:
                if field.name not in ('submit', 'csrf_token', 'contract_id', 'client_id'):
                    try:
                        setattr(g, field.name, field.data if field.data != '' else None)
                    except Exception:
                        pass
            db.session.commit()
            flash('Данные поручителя обновлены', 'success')
            if g.contract_id:
                return redirect(url_for('contract_detail', contract_id=g.contract_id))
            return redirect(url_for('client_detail', client_id=g.client_id))
        guarantor_docs = db.session.execute(
            db.select(Document).where(
                Document.entity_type == 'guarantor', Document.entity_id == g_id
            ).order_by(desc(Document.uploaded_at))
        ).scalars().all()
        return render_template('guarantors/form.html', form=form, contract=contract,
                               client_obj=client_obj, guarantor=g,
                               guarantor_docs=guarantor_docs, title='Редактировать поручителя')

    @app.route('/guarantors/<int:g_id>/delete', methods=['POST'])
    def guarantor_delete(g_id):
        g = db.session.get(Guarantor, g_id) or abort(404)
        cid = g.contract_id
        client_id = g.client_id
        db.session.delete(g)
        db.session.commit()
        flash('Поручитель удалён', 'warning')
        if cid:
            return redirect(url_for('contract_detail', contract_id=cid))
        return redirect(url_for('client_detail', client_id=client_id))

    # ══════════════════════════════════════════════════════════════════════════
    # CONTRACTS
    # ══════════════════════════════════════════════════════════════════════════

    @app.route('/contracts')
    def contracts_list():
        q = request.args.get('q', '').strip()
        status_f = request.args.get('status', '')
        cat_f = request.args.get('category', '')

        stmt = db.select(Contract)
        if q:
            like = f'%{q}%'
            stmt = stmt.where(
                (Contract.contract_number.ilike(like)) |
                (Contract.item_name.ilike(like))
            )
        if status_f:
            stmt = stmt.where(Contract.status == status_f)
        if cat_f:
            stmt = stmt.where(Contract.item_category == cat_f)
        stmt = stmt.order_by(desc(Contract.created_at))
        contracts = db.session.execute(stmt).scalars().all()
        return render_template('contracts/list.html', contracts=contracts,
                               q=q, status_f=status_f, cat_f=cat_f,
                               item_category_choices=ITEM_CATEGORY_CHOICES)

    @app.route('/contracts/new', methods=['GET', 'POST'])
    def contract_new():
        client_id = request.args.get('client_id', type=int)
        client = db.session.get(Client, client_id) if client_id else None
        form = ContractForm()
        if request.method == 'GET':
            form.contract_date.data = date.today()
            form.contract_number.data = generate_contract_number()
            if client_id:
                form.client_id.data = str(client_id)
        if form.validate_on_submit():
            # Check uniqueness of contract_number
            existing = db.session.execute(
                db.select(Contract).where(Contract.contract_number == form.contract_number.data)
            ).scalar_one_or_none()
            if existing:
                form.contract_number.errors.append('Договор с таким номером уже существует')
                clients_all = db.session.execute(
                    db.select(Client).where(Client.status == 'active').order_by(Client.last_name)
                ).scalars().all()
                return render_template('contracts/form.html', form=form, title='Новый договор',
                                       contract=None, client=client, clients_all=clients_all)
            contract = Contract(
                contract_number=form.contract_number.data.strip(),
                client_id=int(form.client_id.data),
                item_name=form.item_name.data,
                item_category=form.item_category.data,
                item_description=form.item_description.data,
                item_serial_number=form.item_serial_number.data,
                item_condition=form.item_condition.data,
                supplier_name=form.supplier_name.data,
                supplier_inn=form.supplier_inn.data,
                supplier_contract_number=form.supplier_contract_number.data,
                cost_price=form.cost_price.data,
                markup_percent=form.markup_percent.data or 0,
                total_price=form.total_price.data,
                down_payment=form.down_payment.data or 0,
                financed_amount=form.financed_amount.data,
                months=form.months.data,
                monthly_payment=form.monthly_payment.data,
                contract_date=form.contract_date.data,
                first_payment_date=form.first_payment_date.data,
                payment_day_of_month=form.payment_day_of_month.data or 5,
                payment_method=form.payment_method.data,
                is_halal=form.is_halal.data,
                halal_certificate_number=form.halal_certificate_number.data,
                sharia_board_approval=form.sharia_board_approval.data,
                sharia_board_note=form.sharia_board_note.data,
                status=form.status.data,
                notes=form.notes.data,
            )
            db.session.add(contract)
            db.session.flush()
            generate_schedule(contract)
            db.session.commit()
            flash(f'Договор {contract.contract_number} создан', 'success')
            return redirect(url_for('contract_detail', contract_id=contract.id))

        clients_all = db.session.execute(
            db.select(Client).where(Client.status == 'active').order_by(Client.last_name)
        ).scalars().all()
        return render_template('contracts/form.html', form=form, title='Новый договор',
                               contract=None, client=client, clients_all=clients_all)

    @app.route('/contracts/<int:contract_id>')
    def contract_detail(contract_id):
        contract = db.session.get(Contract, contract_id) or abort(404)
        schedule = db.session.execute(
            db.select(PaymentSchedule).where(PaymentSchedule.contract_id == contract_id)
            .order_by(asc(PaymentSchedule.installment_num))
        ).scalars().all()
        payments = db.session.execute(
            db.select(Payment).where(Payment.contract_id == contract_id)
            .order_by(desc(Payment.payment_date))
        ).scalars().all()
        guarantors = db.session.execute(
            db.select(Guarantor).where(Guarantor.contract_id == contract_id)
        ).scalars().all()
        documents = db.session.execute(
            db.select(Document).where(
                Document.entity_type == 'contract', Document.entity_id == contract_id
            ).order_by(desc(Document.uploaded_at))
        ).scalars().all()
        schedule_json = [
            {
                'installment_num': s.installment_num,
                'due_date': s.due_date.isoformat() if s.due_date else None,
                'amount': float(s.amount),
                'paid_amount': float(s.paid_amount or 0),
                'status': s.status,
            }
            for s in schedule
        ]
        return render_template('contracts/detail.html',
                               contract=contract, schedule=schedule,
                               schedule_json=schedule_json,
                               payments=payments, guarantors=guarantors,
                               documents=documents, today=date.today())

    @app.route('/contracts/<int:contract_id>/edit', methods=['GET', 'POST'])
    def contract_edit(contract_id):
        contract = db.session.get(Contract, contract_id) or abort(404)
        form = ContractForm(obj=contract)
        if request.method == 'GET':
            form.client_id.data = str(contract.client_id)
        if form.validate_on_submit():
            # Check uniqueness only if number changed
            new_num = form.contract_number.data.strip()
            if new_num != contract.contract_number:
                existing = db.session.execute(
                    db.select(Contract).where(Contract.contract_number == new_num)
                ).scalar_one_or_none()
                if existing:
                    form.contract_number.errors.append('Договор с таким номером уже существует')
                    clients_all = db.session.execute(
                        db.select(Client).where(Client.status == 'active').order_by(Client.last_name)
                    ).scalars().all()
                    return render_template('contracts/form.html', form=form,
                                           title='Редактировать договор',
                                           contract=contract, client=contract.client,
                                           clients_all=clients_all)
            contract.contract_number = new_num
            contract.item_name = form.item_name.data
            contract.item_category = form.item_category.data
            contract.item_description = form.item_description.data
            contract.item_serial_number = form.item_serial_number.data
            contract.item_condition = form.item_condition.data
            contract.supplier_name = form.supplier_name.data
            contract.supplier_inn = form.supplier_inn.data
            contract.supplier_contract_number = form.supplier_contract_number.data
            contract.cost_price = form.cost_price.data
            contract.markup_percent = form.markup_percent.data or 0
            contract.total_price = form.total_price.data
            contract.down_payment = form.down_payment.data or 0
            contract.financed_amount = form.financed_amount.data
            contract.months = form.months.data
            contract.monthly_payment = form.monthly_payment.data
            contract.contract_date = form.contract_date.data
            contract.first_payment_date = form.first_payment_date.data
            contract.payment_day_of_month = form.payment_day_of_month.data or 5
            contract.payment_method = form.payment_method.data
            contract.is_halal = form.is_halal.data
            contract.halal_certificate_number = form.halal_certificate_number.data
            contract.sharia_board_approval = form.sharia_board_approval.data
            contract.sharia_board_note = form.sharia_board_note.data
            contract.status = form.status.data
            contract.notes = form.notes.data
            contract.updated_at = datetime.utcnow()
            generate_schedule(contract)
            db.session.commit()
            flash('Договор обновлён', 'success')
            return redirect(url_for('contract_detail', contract_id=contract.id))

        clients_all = db.session.execute(
            db.select(Client).order_by(Client.last_name)
        ).scalars().all()
        return render_template('contracts/form.html', form=form, title='Редактировать договор',
                               contract=contract, client=contract.client, clients_all=clients_all)

    @app.route('/contracts/<int:contract_id>/close', methods=['POST'])
    def contract_close(contract_id):
        contract = db.session.get(Contract, contract_id) or abort(404)
        contract.status = 'closed'
        contract.early_closure_date = date.today()
        contract.early_closure_reason = request.form.get('reason', 'Досрочное закрытие')
        db.session.commit()
        flash('Договор закрыт', 'success')
        return redirect(url_for('contract_detail', contract_id=contract_id))

    @app.route('/contracts/<int:contract_id>/cancel', methods=['POST'])
    def contract_cancel(contract_id):
        contract = db.session.get(Contract, contract_id) or abort(404)
        contract.status = 'cancelled'
        db.session.commit()
        flash('Договор отменён', 'warning')
        return redirect(url_for('contract_detail', contract_id=contract_id))

    # ══════════════════════════════════════════════════════════════════════════
    # PAYMENTS
    # ══════════════════════════════════════════════════════════════════════════

    @app.route('/payments')
    def payments_list():
        q = request.args.get('q', '').strip()
        date_from = request.args.get('date_from')
        date_to = request.args.get('date_to')

        stmt = db.select(Payment).order_by(desc(Payment.payment_date))
        if q:
            stmt = stmt.join(Contract).where(
                Contract.contract_number.ilike(f'%{q}%')
            )
        if date_from:
            try:
                stmt = stmt.where(Payment.payment_date >= date.fromisoformat(date_from))
            except Exception:
                pass
        if date_to:
            try:
                stmt = stmt.where(Payment.payment_date <= date.fromisoformat(date_to))
            except Exception:
                pass
        payments = db.session.execute(stmt).scalars().all()

        today = date.today()
        today_total = db.session.execute(
            db.select(func.coalesce(func.sum(Payment.amount), 0)).where(
                Payment.payment_date == today)
        ).scalar() or 0
        month_total = db.session.execute(
            db.select(func.coalesce(func.sum(Payment.amount), 0)).where(
                extract('year', Payment.payment_date) == today.year,
                extract('month', Payment.payment_date) == today.month)
        ).scalar() or 0
        all_total = db.session.execute(
            db.select(func.coalesce(func.sum(Payment.amount), 0))
        ).scalar() or 0

        return render_template('payments/list.html', payments=payments, q=q,
                               date_from=date_from, date_to=date_to,
                               today_total=float(today_total),
                               month_total=float(month_total),
                               all_total=float(all_total))

    @app.route('/payments/new', methods=['GET', 'POST'])
    def payment_new():
        contract_id = request.args.get('contract_id', type=int) or request.form.get('contract_id', type=int)
        schedule_id = request.args.get('schedule_id', type=int)
        contract = db.session.get(Contract, contract_id) if contract_id else None

        form = PaymentForm()
        if request.method == 'GET':
            form.payment_date.data = date.today()
            if contract_id:
                form.contract_id.data = str(contract_id)
            if schedule_id:
                form.schedule_id.data = str(schedule_id)
                sched = db.session.get(PaymentSchedule, schedule_id)
                if sched:
                    form.amount.data = sched.remaining

        if form.validate_on_submit():
            cid = int(form.contract_id.data)
            sid = int(form.schedule_id.data) if form.schedule_id.data else None
            payment = Payment(
                contract_id=cid,
                schedule_id=sid,
                amount=form.amount.data,
                payment_date=form.payment_date.data,
                payment_method=form.payment_method.data,
                receipt_number=form.receipt_number.data,
                received_by=form.received_by.data,
                note=form.note.data,
            )
            db.session.add(payment)

            if sid:
                sched = db.session.get(PaymentSchedule, sid)
                if sched:
                    sched.paid_amount += form.amount.data
                    sched.paid_amount = round(sched.paid_amount, 2)
                    if sched.paid_amount >= sched.amount - 0.01:
                        sched.status = 'paid'
                        sched.paid_date = form.payment_date.data
                    else:
                        sched.status = 'partial'

            # Check if contract fully paid
            c = db.session.get(Contract, cid)
            if c:
                total_paid = db.session.execute(
                    db.select(func.coalesce(func.sum(Payment.amount), 0)).where(Payment.contract_id == cid)
                ).scalar() or 0
                total_paid += form.amount.data
                if total_paid >= c.financed_amount - 0.01:
                    c.status = 'closed'
                    c.early_closure_date = form.payment_date.data

            db.session.commit()
            flash('Платёж зарегистрирован', 'success')
            return redirect(url_for('contract_detail', contract_id=cid))

        pending_schedules = []
        if contract:
            pending_schedules = db.session.execute(
                db.select(PaymentSchedule).where(
                    PaymentSchedule.contract_id == contract_id,
                    PaymentSchedule.status.in_(['pending', 'partial', 'overdue'])
                ).order_by(asc(PaymentSchedule.due_date))
            ).scalars().all()

        return render_template('payments/form.html', form=form,
                               contract=contract, pending_schedules=pending_schedules)

    @app.route('/payments/<int:payment_id>/delete', methods=['POST'])
    def payment_delete(payment_id):
        p = db.session.get(Payment, payment_id) or abort(404)
        cid = p.contract_id
        sid = p.schedule_id
        amount = p.amount

        if sid:
            sched = db.session.get(PaymentSchedule, sid)
            if sched:
                sched.paid_amount = max(0, sched.paid_amount - amount)
                if sched.paid_amount <= 0:
                    sched.status = 'pending' if sched.due_date >= date.today() else 'overdue'
                    sched.paid_date = None
                else:
                    sched.status = 'partial'

        db.session.delete(p)
        db.session.commit()
        flash('Платёж удалён', 'warning')
        return redirect(url_for('contract_detail', contract_id=cid))

    # ══════════════════════════════════════════════════════════════════════════
    # REPORTS
    # ══════════════════════════════════════════════════════════════════════════

    @app.route('/reports')
    def reports():
        return render_template('reports.html')

    @app.route('/api/reports/monthly')
    def api_reports_monthly():
        today = date.today()
        result = []
        for i in range(11, -1, -1):
            d = today - relativedelta(months=i)
            total = db.session.execute(
                db.select(func.coalesce(func.sum(Payment.amount), 0)).where(
                    extract('year', Payment.payment_date) == d.year,
                    extract('month', Payment.payment_date) == d.month,
                )
            ).scalar() or 0
            result.append({'month': d.strftime('%b %Y'), 'amount': float(total)})
        return jsonify(result)

    @app.route('/api/reports/stats')
    def api_reports_stats():
        statuses = ['draft', 'active', 'overdue', 'closed', 'cancelled']
        by_status = {}
        for s in statuses:
            cnt = db.session.execute(
                db.select(func.count(Contract.id)).where(Contract.status == s)
            ).scalar() or 0
            by_status[s] = cnt

        categories = ['appliances', 'electronics', 'furniture', 'auto', 'real_estate', 'education', 'medical', 'other']
        by_category = {}
        for cat in categories:
            amount = db.session.execute(
                db.select(func.coalesce(func.sum(Contract.financed_amount), 0)).where(
                    Contract.item_category == cat
                )
            ).scalar() or 0
            by_category[cat] = float(amount)

        overdue_analysis = db.session.execute(
            db.select(Contract).where(Contract.status == 'overdue').order_by(desc(Contract.created_at))
        ).scalars().all()

        overdue_data = []
        today = date.today()
        for c in overdue_analysis:
            overdue_sched = c.overdue_schedules
            if overdue_sched:
                oldest = min(s.due_date for s in overdue_sched)
                days_overdue = (today - oldest).days
                overdue_amount = sum(s.amount - s.paid_amount for s in overdue_sched)
                overdue_data.append({
                    'id': c.id,
                    'contract_number': c.contract_number,
                    'client': c.client.full_name,
                    'days_overdue': days_overdue,
                    'overdue_amount': round(overdue_amount, 2),
                })

        total_portfolio = db.session.execute(
            db.select(func.coalesce(func.sum(Contract.financed_amount), 0)).where(
                Contract.status.in_(['active', 'overdue'])
            )
        ).scalar() or 0

        total_collected = db.session.execute(
            db.select(func.coalesce(func.sum(Payment.amount), 0))
        ).scalar() or 0

        sharia_compliant = db.session.execute(
            db.select(func.count(Contract.id)).where(
                Contract.is_halal == True,
                Contract.status.in_(['active', 'overdue', 'closed'])
            )
        ).scalar() or 0
        total_contracts = db.session.execute(
            db.select(func.count(Contract.id)).where(Contract.status != 'draft')
        ).scalar() or 1

        return jsonify({
            'by_status': by_status,
            'by_category': by_category,
            'overdue_data': overdue_data,
            'total_portfolio': float(total_portfolio),
            'total_collected': float(total_collected),
            'sharia_compliance_pct': round(sharia_compliant / total_contracts * 100, 1),
        })

    # ══════════════════════════════════════════════════════════════════════════
    # BACKUPS
    # ══════════════════════════════════════════════════════════════════════════

    @app.route('/backups')
    def backups_list():
        backups = db.session.execute(
            db.select(Backup).order_by(desc(Backup.created_at))
        ).scalars().all()
        return render_template('backups.html', backups=backups)

    @app.route('/backups/create', methods=['POST'])
    def backup_create():
        note = request.form.get('note', '')
        try:
            filename = backup_module.create_backup(db, note=note)
            flash(f'Резервная копия создана: {filename}', 'success')
        except Exception as e:
            flash(f'Ошибка создания резервной копии: {e}', 'danger')
        return redirect(url_for('backups_list'))

    @app.route('/backups/<int:backup_id>/download')
    def backup_download(backup_id):
        b = db.session.get(Backup, backup_id) or abort(404)
        backup_dir = os.path.join(app.config['DB_DIR'], 'backups')
        filepath = os.path.join(backup_dir, b.filename)
        if not os.path.exists(filepath):
            abort(404)
        return send_file(filepath, as_attachment=True, download_name=b.filename)

    @app.route('/backups/<int:backup_id>/restore', methods=['POST'])
    def backup_restore(backup_id):
        b = db.session.get(Backup, backup_id) or abort(404)
        try:
            backup_module.restore_backup(b.filename, db)
            flash(f'База данных восстановлена из {b.filename}', 'success')
        except Exception as e:
            flash(f'Ошибка восстановления: {e}', 'danger')
        return redirect(url_for('backups_list'))

    # ══════════════════════════════════════════════════════════════════════════
    # DOCUMENTS
    # ══════════════════════════════════════════════════════════════════════════

    @app.route('/documents/upload', methods=['POST'])
    def document_upload():
        entity_type = request.form.get('entity_type', '')
        entity_id = request.form.get('entity_id', type=int)
        doc_type = request.form.get('doc_type', 'other')
        notes = request.form.get('notes', '').strip()

        if entity_type not in ('client', 'contract', 'guarantor') or not entity_id:
            flash('Неверные параметры загрузки', 'danger')
            return redirect(request.referrer or url_for('dashboard'))

        if 'file' not in request.files:
            flash('Файл не выбран', 'danger')
            return redirect(request.referrer or url_for('dashboard'))

        f = request.files['file']
        if not f or f.filename == '':
            flash('Файл не выбран', 'danger')
            return redirect(request.referrer or url_for('dashboard'))

        if not _allowed_file(f.filename):
            flash('Недопустимый формат файла. Разрешены: изображения, PDF, DOC, XLS', 'danger')
            return redirect(request.referrer or url_for('dashboard'))

        ext = f.filename.rsplit('.', 1)[1].lower()
        stored_name = f'{uuid.uuid4().hex}.{ext}'
        filepath = os.path.join(_get_upload_dir(), stored_name)
        f.save(filepath)

        size_kb = round(os.path.getsize(filepath) / 1024, 1)
        doc = Document(
            entity_type=entity_type,
            entity_id=entity_id,
            doc_type=doc_type,
            filename=stored_name,
            original_name=secure_filename(f.filename),
            file_size_kb=size_kb,
            mime_type=f.content_type or '',
            notes=notes or None,
        )
        db.session.add(doc)
        db.session.commit()
        flash('Документ загружен', 'success')
        return redirect(request.referrer or url_for('dashboard'))

    @app.route('/documents/<int:doc_id>/file')
    def document_download(doc_id):
        doc = db.session.get(Document, doc_id) or abort(404)
        filepath = os.path.join(_get_upload_dir(), doc.filename)
        if not os.path.exists(filepath):
            abort(404)
        return send_file(filepath, download_name=doc.original_name or doc.filename, as_attachment=False)

    @app.route('/documents/<int:doc_id>/delete', methods=['POST'])
    def document_delete(doc_id):
        doc = db.session.get(Document, doc_id) or abort(404)
        filepath = os.path.join(_get_upload_dir(), doc.filename)
        if os.path.exists(filepath):
            os.remove(filepath)
        db.session.delete(doc)
        db.session.commit()
        flash('Документ удалён', 'warning')
        return redirect(request.referrer or url_for('dashboard'))

    @app.route('/backups/<int:backup_id>/delete', methods=['POST'])
    def backup_delete(backup_id):
        b = db.session.get(Backup, backup_id) or abort(404)
        backup_dir = os.path.join(app.config['DB_DIR'], 'backups')
        filepath = os.path.join(backup_dir, b.filename)
        if os.path.exists(filepath):
            os.remove(filepath)
        db.session.delete(b)
        db.session.commit()
        flash('Резервная копия удалена', 'warning')
        return redirect(url_for('backups_list'))

    # ══════════════════════════════════════════════════════════════════════════
    # IMPORT FROM EXCEL
    # ══════════════════════════════════════════════════════════════════════════

    def _parse_date_ru(val):
        """Парсит дату из строки ДД.ММ.ГГГГ или ГГГГ-ММ-ДД."""
        if not val:
            return None
        if isinstance(val, (date, datetime)):
            return val if isinstance(val, date) else val.date()
        s = str(val).strip()
        for fmt in ('%d.%m.%Y', '%Y-%m-%d', '%d/%m/%Y', '%m/%d/%Y'):
            try:
                return datetime.strptime(s, fmt).date()
            except ValueError:
                pass
        return None

    def _str(val):
        return str(val).strip() if val is not None else ''

    def _float(val, default=0.0):
        try:
            return float(str(val).replace(',', '.').replace(' ', '')) if val is not None else default
        except (ValueError, TypeError):
            return default

    def _int(val, default=0):
        try:
            return int(float(str(val))) if val is not None else default
        except (ValueError, TypeError):
            return default

    def _read_excel_rows(file_storage):
        """Читает .xlsx или .csv, возвращает (headers, rows)."""
        fname = file_storage.filename.lower()
        if fname.endswith('.csv'):
            content = file_storage.read().decode('utf-8-sig')
            reader = csv.reader(io.StringIO(content), delimiter=';')
            rows = list(reader)
            if not rows:
                return [], []
            return [h.strip() for h in rows[0]], rows[1:]
        else:
            wb = openpyxl.load_workbook(file_storage, data_only=True)
            ws = wb.active
            all_rows = list(ws.iter_rows(values_only=True))
            if not all_rows:
                return [], []
            headers = [str(h).strip() if h is not None else '' for h in all_rows[0]]
            return headers, [list(r) for r in all_rows[1:]]

    def _make_template(title, columns, example_row, filename):
        """Создаёт xlsx-шаблон и возвращает Response."""
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = title

        hdr_fill = PatternFill('solid', fgColor='1E3A5F')
        hdr_font = Font(color='FFFFFF', bold=True, size=11)
        ex_fill  = PatternFill('solid', fgColor='F0F4F8')
        ex_font  = Font(color='888888', italic=True)

        for col_idx, col_name in enumerate(columns, 1):
            cell = ws.cell(row=1, column=col_idx, value=col_name)
            cell.font = hdr_fill and hdr_font
            cell.fill = hdr_fill
            cell.alignment = Alignment(horizontal='center', wrap_text=True)
            ws.column_dimensions[openpyxl.utils.get_column_letter(col_idx)].width = max(len(col_name) + 4, 14)

        for col_idx, val in enumerate(example_row, 1):
            cell = ws.cell(row=2, column=col_idx, value=val)
            cell.font = ex_font
            cell.fill = ex_fill

        ws.row_dimensions[1].height = 30
        ws.freeze_panes = 'A2'

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return send_file(buf, as_attachment=True, download_name=filename,
                         mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

    @app.route('/import')
    def import_page():
        return render_template('import.html')

    # ── Шаблоны ──────────────────────────────────────────────────────────────

    @app.route('/import/template/clients')
    def import_template_clients():
        cols = [
            'Фамилия*', 'Имя*', 'Отчество', 'Телефон*', 'Телефон 2', 'Email',
            'Дата рождения (ДД.ММ.ГГГГ)', 'Пол (М/Ж)', 'ИНН', 'СНИЛС',
            'Серия паспорта', 'Номер паспорта', 'Кем выдан', 'Дата выдачи (ДД.ММ.ГГГГ)',
            'Адрес прописка', 'Адрес факт.', 'Занятость (employed/self_employed/business_owner/unemployed)',
            'Работодатель', 'Должность', 'Стаж (мес)', 'Доход в мес (руб)',
            'Семейное положение (single/married/divorced/widowed)', 'Дети',
            'Кредитный балл (0-999)', 'Примечания',
        ]
        example = [
            'Иванов', 'Иван', 'Иванович', '+7-900-000-00-01', '', 'ivan@mail.ru',
            '15.03.1985', 'М', '123456789012', '', '1234', '567890', 'МВД', '01.01.2015',
            'г. Москва, ул. Ленина, д.1', '', 'employed',
            'ООО Ромашка', 'Менеджер', '36', '80000',
            'married', '2', '650', '',
        ]
        return _make_template('Клиенты', cols, example, 'шаблон_клиенты.xlsx')

    @app.route('/import/template/contracts')
    def import_template_contracts():
        cols = [
            'Телефон клиента*', 'Номер договора', 'Дата договора* (ДД.ММ.ГГГГ)',
            'Наименование товара*', 'Категория (appliances/electronics/furniture/auto/real_estate/education/medical/other)',
            'Себестоимость*', 'Наценка %*', 'Первый взнос', 'Срок (мес)*',
            'Дата первого платежа* (ДД.ММ.ГГГГ)', 'День платежа (1-28)',
            'Способ оплаты (cash/bank_transfer/card/online)', 'Примечания',
        ]
        example = [
            '+7-900-000-00-01', 'МУР-2024-001', '01.06.2024',
            'Холодильник Samsung', 'appliances',
            '45000', '20', '5000', '12',
            '01.07.2024', '1',
            'cash', '',
        ]
        return _make_template('Договоры', cols, example, 'шаблон_договоры.xlsx')

    @app.route('/import/template/payments')
    def import_template_payments():
        cols = [
            'Номер договора*', 'Дата платежа* (ДД.ММ.ГГГГ)', 'Сумма*',
            'Способ оплаты (cash/bank_transfer/card/online)',
            'Квитанция №', 'Принял',
        ]
        example = [
            'МУР-2024-001', '01.08.2024', '4167',
            'cash', 'КВ-001', 'Администратор',
        ]
        return _make_template('Платежи', cols, example, 'шаблон_платежи.xlsx')

    # ── Импорт клиентов ───────────────────────────────────────────────────────

    @app.route('/import/clients', methods=['POST'])
    def import_clients():
        f = request.files.get('file')
        if not f or f.filename == '':
            flash('Файл не выбран', 'danger')
            return redirect(url_for('import_page'))

        try:
            headers, rows = _read_excel_rows(f)
        except Exception as e:
            flash(f'Ошибка чтения файла: {e}', 'danger')
            return redirect(url_for('import_page'))

        # Нормализуем заголовки
        h = [x.replace('*', '').strip().lower() for x in headers]

        def col(name_variants):
            for v in name_variants:
                for i, hdr in enumerate(h):
                    if v.lower() in hdr:
                        return i
            return None

        i_last   = col(['фамилия', 'lastname', 'last_name'])
        i_first  = col(['имя', 'firstname', 'first_name'])
        i_mid    = col(['отчество', 'middlename', 'middle'])
        i_phone  = col(['телефон', 'phone', 'тел'])
        i_phone2 = col(['телефон 2', 'phone2', 'phone 2', 'тел 2'])
        i_email  = col(['email', 'почта'])
        i_bdate  = col(['дата рожд', 'birth'])
        i_gender = col(['пол', 'gender'])
        i_inn    = col(['инн', 'inn'])
        i_snils  = col(['снилс', 'snils'])
        i_ps     = col(['серия паспорта', 'passport_series', 'серия'])
        i_pn     = col(['номер паспорта', 'passport_number'])
        i_pb     = col(['кем выдан', 'issued_by'])
        i_pd     = col(['дата выдачи', 'issued_date'])
        i_addr   = col(['адрес прописка', 'адрес регистр', 'address_reg'])
        i_addr2  = col(['адрес факт', 'address_act', 'фактический'])
        i_emp    = col(['занятость', 'employment'])
        i_emplr  = col(['работодатель', 'employer'])
        i_pos    = col(['должность', 'position'])
        i_exp    = col(['стаж', 'experience'])
        i_inc    = col(['доход', 'income'])
        i_marr   = col(['семейное', 'marital'])
        i_chld   = col(['дети', 'children'])
        i_score  = col(['балл', 'score', 'кредитный'])
        i_notes  = col(['примечания', 'notes', 'note'])

        if i_last is None or i_first is None or i_phone is None:
            flash('Файл должен содержать колонки: Фамилия, Имя, Телефон', 'danger')
            return redirect(url_for('import_page'))

        def cell(row, idx):
            return _str(row[idx]) if idx is not None and idx < len(row) else ''

        gender_map = {'м': 'male', 'м.': 'male', 'муж': 'male', 'male': 'male',
                      'ж': 'female', 'ж.': 'female', 'жен': 'female', 'female': 'female'}

        created = skipped = errors = 0
        error_msgs = []

        for row_num, row in enumerate(rows, 2):
            if not any(row):
                continue
            last  = cell(row, i_last)
            first = cell(row, i_first)
            phone = cell(row, i_phone)
            if not last or not first or not phone:
                errors += 1
                error_msgs.append(f'Строка {row_num}: пропущены обязательные поля (Фамилия/Имя/Телефон)')
                continue

            # Проверка дубликата по телефону
            existing = db.session.execute(
                db.select(Client).where(Client.phone == phone)
            ).scalar_one_or_none()
            if existing:
                skipped += 1
                continue

            g_raw = gender_map.get(cell(row, i_gender).lower(), None)
            emp_val = cell(row, i_emp)
            emp_map = {
                'наёмный': 'employed', 'наемный': 'employed', 'employed': 'employed',
                'самозанятый': 'self_employed', 'self_employed': 'self_employed',
                'бизнес': 'business_owner', 'business_owner': 'business_owner',
                'безработный': 'unemployed', 'unemployed': 'unemployed',
            }
            emp_norm = emp_map.get(emp_val.lower(), emp_val if emp_val else None)

            marr_map = {'single': 'single', 'холост': 'single', 'не женат': 'single',
                        'married': 'married', 'женат': 'married', 'замужем': 'married',
                        'divorced': 'divorced', 'разведён': 'divorced', 'разведена': 'divorced',
                        'widowed': 'widowed', 'вдовец': 'widowed', 'вдова': 'widowed'}
            marr_raw = cell(row, i_marr)
            marr_norm = marr_map.get(marr_raw.lower(), None)

            client = Client(
                last_name=last, first_name=first,
                middle_name=cell(row, i_mid) or None,
                phone=phone,
                phone2=cell(row, i_phone2) or None,
                email=cell(row, i_email) or None,
                birth_date=_parse_date_ru(cell(row, i_bdate)),
                gender=g_raw,
                inn=cell(row, i_inn) or None,
                snils=cell(row, i_snils) or None,
                passport_series=cell(row, i_ps) or None,
                passport_number=cell(row, i_pn) or None,
                passport_issued_by=cell(row, i_pb) or None,
                passport_issued_date=_parse_date_ru(cell(row, i_pd)),
                address_registration=cell(row, i_addr) or None,
                address_actual=cell(row, i_addr2) or None,
                employment_type=emp_norm,
                employer_name=cell(row, i_emplr) or None,
                position=cell(row, i_pos) or None,
                work_experience_months=_int(cell(row, i_exp)) or None,
                monthly_income=_float(cell(row, i_inc)),
                marital_status=marr_norm,
                children_count=_int(cell(row, i_chld)),
                credit_score=_int(cell(row, i_score)),
                notes=cell(row, i_notes) or None,
                status='active',
            )
            db.session.add(client)
            created += 1

        try:
            db.session.commit()
        except Exception as e:
            db.session.rollback()
            flash(f'Ошибка сохранения: {e}', 'danger')
            return redirect(url_for('import_page'))

        msg = f'Клиенты: добавлено {created}'
        if skipped:
            msg += f', пропущено дубликатов {skipped}'
        if errors:
            msg += f', ошибок {errors}'
        flash(msg, 'success' if created > 0 else 'warning')
        if error_msgs:
            flash('Ошибки строк: ' + '; '.join(error_msgs[:5]), 'warning')
        return redirect(url_for('clients_list'))

    # ── Импорт договоров ──────────────────────────────────────────────────────

    @app.route('/import/contracts', methods=['POST'])
    def import_contracts():
        f = request.files.get('file')
        if not f or f.filename == '':
            flash('Файл не выбран', 'danger')
            return redirect(url_for('import_page'))

        try:
            headers, rows = _read_excel_rows(f)
        except Exception as e:
            flash(f'Ошибка чтения файла: {e}', 'danger')
            return redirect(url_for('import_page'))

        h = [x.replace('*', '').strip().lower() for x in headers]

        def col(variants):
            for v in variants:
                for i, hdr in enumerate(h):
                    if v.lower() in hdr:
                        return i
            return None

        def cell(row, idx):
            return _str(row[idx]) if idx is not None and idx < len(row) else ''

        i_phone   = col(['телефон', 'phone'])
        i_num     = col(['номер договора', 'contract_number', 'номер'])
        i_cdate   = col(['дата договора', 'contract_date'])
        i_item    = col(['наименование', 'товар', 'item_name'])
        i_cat     = col(['категория', 'category'])
        i_cost    = col(['себестоимость', 'cost'])
        i_markup  = col(['наценка', 'markup'])
        i_down    = col(['первый взнос', 'down_payment', 'взнос'])
        i_months  = col(['срок', 'months'])
        i_fpdate  = col(['дата первого', 'first_payment'])
        i_pday    = col(['день платежа', 'payment_day'])
        i_method  = col(['способ', 'method'])
        i_notes   = col(['примечания', 'notes'])

        if i_phone is None or i_cost is None or i_months is None:
            flash('Файл должен содержать: Телефон клиента, Себестоимость, Срок', 'danger')
            return redirect(url_for('import_page'))

        created = skipped = errors = 0
        error_msgs = []

        for row_num, row in enumerate(rows, 2):
            if not any(row):
                continue
            phone = cell(row, i_phone)
            if not phone:
                continue

            client = db.session.execute(
                db.select(Client).where(Client.phone == phone)
            ).scalar_one_or_none()
            if not client:
                errors += 1
                error_msgs.append(f'Строка {row_num}: клиент с телефоном {phone} не найден')
                continue

            cost    = _float(cell(row, i_cost))
            markup  = _float(cell(row, i_markup), 0.0)
            down    = _float(cell(row, i_down),   0.0)
            months  = _int(cell(row, i_months), 12)
            total   = round(cost * (1 + markup / 100), 2)
            financed = max(0, total - down)
            monthly = round(financed / months, 2) if months else 0

            cdate  = _parse_date_ru(cell(row, i_cdate))  or date.today()
            fpdate = _parse_date_ru(cell(row, i_fpdate)) or date(cdate.year, cdate.month + 1 if cdate.month < 12 else 1, 1)

            num = cell(row, i_num) or generate_contract_number()

            # Дубликат по номеру
            if db.session.execute(db.select(Contract).where(Contract.contract_number == num)).scalar_one_or_none():
                skipped += 1
                continue

            method_map = {'наличные': 'cash', 'cash': 'cash',
                          'банк': 'bank_transfer', 'bank_transfer': 'bank_transfer',
                          'карта': 'card', 'card': 'card',
                          'онлайн': 'online', 'online': 'online'}
            method_raw = cell(row, i_method).lower()
            method = method_map.get(method_raw, 'cash')

            cat_map = {'техника': 'appliances', 'appliances': 'appliances',
                       'электроника': 'electronics', 'electronics': 'electronics',
                       'мебель': 'furniture', 'furniture': 'furniture',
                       'авто': 'auto', 'auto': 'auto',
                       'недвижимость': 'real_estate', 'real_estate': 'real_estate',
                       'образование': 'education', 'education': 'education',
                       'медицина': 'medical', 'medical': 'medical'}
            cat_raw = cell(row, i_cat).lower()
            cat = cat_map.get(cat_raw, 'other')

            pday = _int(cell(row, i_pday), fpdate.day)

            contract = Contract(
                client_id=client.id,
                contract_number=num,
                contract_date=cdate,
                first_payment_date=fpdate,
                item_name=cell(row, i_item) or 'Товар',
                item_category=cat,
                cost_price=cost,
                markup_percent=markup,
                total_price=total,
                down_payment=down,
                financed_amount=financed,
                months=months,
                monthly_payment=monthly,
                payment_day_of_month=pday,
                payment_method=method,
                notes=cell(row, i_notes) or None,
                status='active',
                is_halal=True,
            )
            db.session.add(contract)
            db.session.flush()

            # Генерируем график
            generate_schedule(contract)
            created += 1

        try:
            db.session.commit()
            # Обновляем просроченные после импорта
            update_overdue_statuses()
        except Exception as e:
            db.session.rollback()
            flash(f'Ошибка сохранения: {e}', 'danger')
            return redirect(url_for('import_page'))

        msg = f'Договоры: добавлено {created}'
        if skipped: msg += f', пропущено {skipped}'
        if errors:  msg += f', ошибок {errors}'
        flash(msg, 'success' if created > 0 else 'warning')
        if error_msgs:
            flash('Ошибки: ' + '; '.join(error_msgs[:5]), 'warning')
        return redirect(url_for('contracts_list'))

    # ── Импорт платежей (история) ─────────────────────────────────────────────

    @app.route('/import/payments', methods=['POST'])
    def import_payments():
        f = request.files.get('file')
        if not f or f.filename == '':
            flash('Файл не выбран', 'danger')
            return redirect(url_for('import_page'))

        try:
            headers, rows = _read_excel_rows(f)
        except Exception as e:
            flash(f'Ошибка чтения файла: {e}', 'danger')
            return redirect(url_for('import_page'))

        h = [x.replace('*', '').strip().lower() for x in headers]

        def col(variants):
            for v in variants:
                for i, hdr in enumerate(h):
                    if v.lower() in hdr:
                        return i
            return None

        def cell(row, idx):
            return _str(row[idx]) if idx is not None and idx < len(row) else ''

        i_num    = col(['номер договора', 'contract', 'договор'])
        i_date   = col(['дата платежа', 'payment_date', 'дата'])
        i_amount = col(['сумма', 'amount'])
        i_method = col(['способ', 'method'])
        i_rcpt   = col(['квитанция', 'receipt'])
        i_by     = col(['принял', 'received_by'])

        if i_num is None or i_date is None or i_amount is None:
            flash('Файл должен содержать: Номер договора, Дата платежа, Сумма', 'danger')
            return redirect(url_for('import_page'))

        method_map = {'наличные': 'cash', 'cash': 'cash',
                      'банк': 'bank_transfer', 'bank_transfer': 'bank_transfer',
                      'карта': 'card', 'card': 'card',
                      'онлайн': 'online', 'online': 'online'}

        created = errors = 0
        error_msgs = []

        for row_num, row in enumerate(rows, 2):
            if not any(row):
                continue
            num    = cell(row, i_num)
            pdate  = _parse_date_ru(cell(row, i_date))
            amount = _float(cell(row, i_amount))

            if not num or not pdate or amount <= 0:
                errors += 1
                error_msgs.append(f'Строка {row_num}: номер договора, дата или сумма не указаны')
                continue

            contract = db.session.execute(
                db.select(Contract).where(Contract.contract_number == num)
            ).scalar_one_or_none()
            if not contract:
                errors += 1
                error_msgs.append(f'Строка {row_num}: договор {num!r} не найден')
                continue

            method_raw = cell(row, i_method).lower()
            method = method_map.get(method_raw, 'cash')

            payment = Payment(
                contract_id=contract.id,
                amount=amount,
                payment_date=pdate,
                payment_method=method,
                receipt_number=cell(row, i_rcpt) or None,
                received_by=cell(row, i_by) or None,
            )
            db.session.add(payment)

            # Привязываем к ближайшей неоплаченной строке графика
            pending = db.session.execute(
                db.select(PaymentSchedule)
                .where(
                    PaymentSchedule.contract_id == contract.id,
                    PaymentSchedule.status.in_(['pending', 'overdue', 'partial']),
                )
                .order_by(asc(PaymentSchedule.due_date))
            ).scalars().first()

            if pending:
                new_paid = (pending.paid_amount or 0) + amount
                if new_paid >= pending.amount - 0.01:
                    pending.paid_amount = pending.amount
                    pending.status = 'paid'
                    pending.paid_date = pdate
                else:
                    pending.paid_amount = new_paid
                    pending.status = 'partial'
                payment.schedule_id = pending.id

            created += 1

        try:
            db.session.commit()
            update_overdue_statuses()
        except Exception as e:
            db.session.rollback()
            flash(f'Ошибка сохранения: {e}', 'danger')
            return redirect(url_for('import_page'))

        msg = f'Платежи: добавлено {created}'
        if errors: msg += f', ошибок {errors}'
        flash(msg, 'success' if created > 0 else 'warning')
        if error_msgs:
            flash('Ошибки: ' + '; '.join(error_msgs[:5]), 'warning')
        return redirect(url_for('payments_list'))

    # ══════════════════════════════════════════════════════════════════════════
    # CALCULATOR
    # ══════════════════════════════════════════════════════════════════════════

    @app.route('/calculator')
    def calculator():
        return render_template('calculator.html')

    # ══════════════════════════════════════════════════════════════════════════
    # CONTRACT DOCUMENT GENERATION
    # ══════════════════════════════════════════════════════════════════════════

    def _contracts_dir():
        d = os.path.join(app.config['DB_DIR'], 'contracts')
        os.makedirs(d, exist_ok=True)
        return d

    def _template_path():
        user_tpl = os.path.join(app.config['DB_DIR'], 'contract_template.docx')
        if os.path.exists(user_tpl):
            return user_tpl
        # Bundled default template (shipped with the app)
        base = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))
        return os.path.join(base, 'default_contract_template.docx')

    MONTHS_RU = ['января','февраля','марта','апреля','мая','июня',
                 'июля','августа','сентября','октября','ноября','декабря']

    def _fmt_date_ru(d):
        if not d:
            return ''
        return f"{d.day} {MONTHS_RU[d.month-1]} {d.year} г."

    def _replace_in_paragraph(para, repl):
        """Заменяет метки в параграфе, сохраняя форматирование первого run-а."""
        full = ''.join(r.text for r in para.runs)
        if not any(k in full for k in repl):
            return
        for k, v in repl.items():
            full = full.replace(k, str(v) if v is not None else '')
        if para.runs:
            para.runs[0].text = full
            for r in para.runs[1:]:
                r.text = ''

    def _build_replacements(contract):
        client = contract.client
        cdate  = contract.contract_date
        replacements = {
            # Договор
            '{{НОМЕР_ДОГОВОРА}}':        contract.contract_number or '',
            '{{ДАТА_ДОГОВОРА}}':         _fmt_date_ru(cdate),
            '{{ЧИСЛО}}':                 str(cdate.day) if cdate else '',
            '{{МЕСЯЦ}}':                 MONTHS_RU[cdate.month-1].capitalize() if cdate else '',
            '{{ГОД}}':                   str(cdate.year) if cdate else '',
            # Покупатель
            '{{ФИО_ПОКУПАТЕЛЯ}}':        client.full_name,
            '{{ДАТА_РОЖДЕНИЯ}}':         client.birth_date.strftime('%d.%m.%Y') if client.birth_date else '',
            '{{ПАСПОРТ_СЕРИЯ}}':         client.passport_series or '',
            '{{ПАСПОРТ_НОМЕР}}':         client.passport_number or '',
            '{{ПАСПОРТ_ВЫДАН}}':         client.passport_issued_by or '',
            '{{ПАСПОРТ_ДАТА}}':          client.passport_issued_date.strftime('%d.%m.%Y') if client.passport_issued_date else '',
            '{{ПАСПОРТ_КОД}}':           '',
            '{{АДРЕС}}':                 client.address_registration or client.address_actual or '',
            '{{ТЕЛЕФОН}}':               client.phone or '',
            # Товар
            '{{ТОВАР}}':                 contract.item_name or '',
            '{{КАТЕГОРИЯ}}':             contract.item_category or '',
            '{{СЕРИЙНЫЙ_НОМЕР}}':        contract.item_serial_number or '',
            '{{ПОСТАВЩИК}}':             contract.supplier_name or '',
            # Финансы
            '{{СЕБЕСТОИМОСТЬ}}':         f"{contract.cost_price:,.0f}".replace(',', ' '),
            '{{НАЦЕНКА_ПРОЦЕНТ}}':       str(contract.markup_percent),
            '{{ИТОГОВАЯ_СУММА}}':        f"{contract.total_price:,.0f}".replace(',', ' '),
            '{{ПЕРВЫЙ_ВЗНОС}}':          f"{contract.down_payment:,.0f}".replace(',', ' '),
            '{{ФИНАНСИРУЕТСЯ}}':         f"{contract.financed_amount:,.0f}".replace(',', ' '),
            '{{СРОК_МЕСЯЦЕВ}}':          str(contract.months),
            '{{ЕЖЕМЕСЯЧНЫЙ_ПЛАТЕЖ}}':    f"{contract.monthly_payment:,.0f}".replace(',', ' '),
            '{{ДАТА_ПЕРВОГО_ПЛАТЕЖА}}':  _fmt_date_ru(contract.first_payment_date),
            '{{ДАТА_ПОСЛЕДНЕГО_ПЛАТЕЖА}}': _fmt_date_ru(contract.last_payment_date),
            '{{СПОСОБ_ОПЛАТЫ}}':         {'cash':'Наличные','bank_transfer':'Банковский перевод',
                                           'card':'Карта','online':'Онлайн'}.get(contract.payment_method,''),
        }
        return replacements

    def _generate_doc(contract):
        """Генерирует .docx из шаблона и сохраняет в contracts/."""
        tpl = _template_path()
        if not os.path.exists(tpl):
            return None, 'Шаблон договора не загружен'
        try:
            doc  = DocxDocument(tpl)
            repl = _build_replacements(contract)

            for para in doc.paragraphs:
                _replace_in_paragraph(para, repl)
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            _replace_in_paragraph(para, repl)

            safe_num = contract.contract_number.replace('/', '-').replace('\\', '-')
            filename = f"{safe_num}.docx"
            filepath = os.path.join(_contracts_dir(), filename)
            doc.save(filepath)
            return filename, None
        except Exception as e:
            return None, str(e)

    @app.route('/settings/doc-template', methods=['GET', 'POST'])
    def doc_template_settings():
        tpl = _template_path()
        exists = os.path.exists(tpl)
        if request.method == 'POST':
            f = request.files.get('template')
            if f and f.filename.endswith('.docx'):
                f.save(tpl)
                flash('Шаблон договора загружен', 'success')
            else:
                flash('Загрузите файл .docx', 'danger')
            return redirect(url_for('doc_template_settings'))
        return render_template('doc_template.html', exists=exists,
                               tpl_name='contract_template.docx' if exists else None)

    @app.route('/contracts/<int:contract_id>/generate-doc', methods=['POST'])
    def contract_generate_doc(contract_id):
        contract = db.session.get(Contract, contract_id) or abort(404)
        filename, err = _generate_doc(contract)
        if err:
            flash(f'Ошибка генерации: {err}', 'danger')
        else:
            flash('Договор сформирован', 'success')
        return redirect(url_for('contract_detail', contract_id=contract_id))

    @app.route('/contracts/<int:contract_id>/download-doc')
    def contract_download_doc(contract_id):
        contract = db.session.get(Contract, contract_id) or abort(404)
        safe_num = contract.contract_number.replace('/', '-').replace('\\', '-')
        filename = f"{safe_num}.docx"
        filepath = os.path.join(_contracts_dir(), filename)
        if not os.path.exists(filepath):
            # Попробуем сгенерировать автоматически
            filename, err = _generate_doc(contract)
            if err:
                flash(f'Нет файла договора: {err}', 'danger')
                return redirect(url_for('contract_detail', contract_id=contract_id))
            filepath = os.path.join(_contracts_dir(), filename)
        return send_file(filepath, as_attachment=False,
                         download_name=f"Договор {contract.contract_number}.docx")

    # ══════════════════════════════════════════════════════════════════════════
    # SEED DATA
    # ══════════════════════════════════════════════════════════════════════════

    def seed_demo_data():
        if db.session.execute(db.select(func.count(Client.id))).scalar() > 0:
            return

        c1 = Client(
            last_name='Алиев', first_name='Руслан', middle_name='Магомедович',
            birth_date=date(1985, 3, 15), gender='male', citizenship='Россия',
            passport_series='2214', passport_number='345678',
            passport_issued_by='УФМС России по г. Москве', passport_issued_date=date(2015, 6, 1),
            inn='7720123456', snils='123-456-789 01',
            phone='+7-905-123-45-67', email='alieff@example.com',
            address_registration='г. Москва, ул. Ленина, д. 10, кв. 5',
            address_actual='г. Москва, ул. Садовая, д. 3, кв. 12',
            employer_name='ООО "РосТехник"', position='Инженер',
            employment_type='employed', work_experience_months=60,
            monthly_income=85000, income_confirmed=True,
            marital_status='married', children_count=2,
            education='university', own_property=True, own_car=True,
            credit_history_status='good', credit_score=750, status='active',
        )

        c2 = Client(
            last_name='Мусаева', first_name='Заира', middle_name='Исмаиловна',
            birth_date=date(1990, 7, 22), gender='female', citizenship='Россия',
            passport_series='0512', passport_number='987654',
            passport_issued_by='УФМС России по Республике Дагестан',
            passport_issued_date=date(2018, 9, 10),
            inn='0501234567', phone='+7-928-456-78-90',
            address_registration='г. Махачкала, ул. Гагарина, д. 25, кв. 8',
            address_actual='г. Махачкала, ул. Гагарина, д. 25, кв. 8',
            employer_name='ИП Мусаева З.И.', position='Предприниматель',
            employment_type='self_employed', work_experience_months=36,
            monthly_income=55000, income_confirmed=False,
            marital_status='married', children_count=1,
            education='college', credit_history_status='satisfactory',
            credit_score=580, status='active',
        )

        c3 = Client(
            last_name='Хасанов', first_name='Тимур', middle_name='Рустамович',
            birth_date=date(1978, 11, 5), gender='male', citizenship='Россия',
            passport_series='0308', passport_number='112233',
            passport_issued_by='ОВД г. Казани', passport_issued_date=date(2010, 4, 20),
            inn='1650123456', phone='+7-917-789-01-23',
            address_registration='г. Казань, ул. Баумана, д. 15, кв. 22',
            address_actual='г. Казань, ул. Баумана, д. 15, кв. 22',
            employer_name='АО "КазТрансСтрой"', position='Директор',
            employment_type='business_owner', work_experience_months=120,
            monthly_income=150000, income_confirmed=True,
            marital_status='married', children_count=3,
            education='university', own_property=True, own_car=True,
            credit_history_status='good', credit_score=820, status='active',
        )

        db.session.add_all([c1, c2, c3])
        db.session.flush()

        # Contract 1 – active
        today = date.today()
        con1 = Contract(
            contract_number='МУР-2024-001',
            client_id=c1.id,
            item_name='Холодильник Samsung RF65A967FSR',
            item_category='appliances',
            item_condition='new',
            supplier_name='М.Видео', supplier_inn='7812345678',
            cost_price=85000, markup_percent=15.0, total_price=97750,
            down_payment=10000, financed_amount=87750,
            months=12, monthly_payment=7312.5,
            contract_date=today - timedelta(days=90),
            first_payment_date=today - timedelta(days=60),
            payment_day_of_month=5, payment_method='cash',
            is_halal=True, sharia_board_approval=True,
            status='active',
        )

        # Contract 2 – overdue
        con2 = Contract(
            contract_number='МУР-2024-002',
            client_id=c2.id,
            item_name='Ноутбук ASUS VivoBook 15',
            item_category='electronics',
            item_condition='new',
            supplier_name='DNS', supplier_inn='2509987654',
            cost_price=55000, markup_percent=10.0, total_price=60500,
            down_payment=5000, financed_amount=55500,
            months=6, monthly_payment=9250,
            contract_date=today - timedelta(days=180),
            first_payment_date=today - timedelta(days=150),
            payment_day_of_month=10, payment_method='bank_transfer',
            is_halal=True, sharia_board_approval=True,
            status='active',
        )

        # Contract 3 – draft
        con3 = Contract(
            contract_number='МУР-2024-003',
            client_id=c3.id,
            item_name='Автомобиль Toyota Camry 2023',
            item_category='auto',
            item_condition='new',
            supplier_name='Тойота-Казань', supplier_inn='1650543210',
            cost_price=2500000, markup_percent=12.0, total_price=2800000,
            down_payment=500000, financed_amount=2300000,
            months=36, monthly_payment=63888.89,
            contract_date=today,
            first_payment_date=today + timedelta(days=30),
            payment_day_of_month=15, payment_method='bank_transfer',
            is_halal=True, sharia_board_approval=False,
            status='draft',
        )

        db.session.add_all([con1, con2, con3])
        db.session.flush()

        # Generate schedules
        generate_schedule(con1)
        generate_schedule(con2)
        generate_schedule(con3)

        # Add some payments for contract 1
        sched1 = db.session.execute(
            db.select(PaymentSchedule).where(
                PaymentSchedule.contract_id == con1.id,
                PaymentSchedule.installment_num == 1
            )
        ).scalar_one_or_none()
        if sched1:
            p1 = Payment(
                contract_id=con1.id, schedule_id=sched1.id,
                amount=7312.5, payment_date=sched1.due_date,
                payment_method='cash', receipt_number='ЧЕК-001',
                received_by='Администратор',
            )
            sched1.paid_amount = 7312.5
            sched1.status = 'paid'
            sched1.paid_date = sched1.due_date
            db.session.add(p1)

        # Guarantor for contract 1
        g1 = Guarantor(
            contract_id=con1.id,
            last_name='Алиев', first_name='Магомед', middle_name='Алиевич',
            phone='+7-905-000-11-22', relationship='parent',
            passport_series='2210', passport_number='654321',
            passport_issued_by='УФМС г. Москвы',
            passport_issued_date=date(2010, 1, 1),
            address_registration='г. Москва, ул. Ленина, д. 10, кв. 3',
            employer_name='Пенсионер', monthly_income=25000,
            guarantor_type='personal',
        )
        db.session.add(g1)
        db.session.commit()

    # Init DB
    with app.app_context():
        db.create_all()
        # Best-effort migration for existing databases
        try:
            from sqlalchemy import text, inspect as sa_inspect
            insp = sa_inspect(db.engine)
            g_cols = {c['name'] for c in insp.get_columns('guarantors')} if 'guarantors' in insp.get_table_names() else set()
            new_g_cols = [
                ('birth_date', 'DATE'), ('gender', 'VARCHAR(10)'), ('phone2', 'VARCHAR(30)'),
                ('email', 'VARCHAR(150)'), ('inn', 'VARCHAR(20)'), ('snils', 'VARCHAR(20)'),
                ('address_actual', 'TEXT'), ('employer_phone', 'VARCHAR(30)'),
                ('position', 'VARCHAR(150)'), ('employment_type', 'VARCHAR(30)'),
                ('work_experience_months', 'INTEGER'), ('created_at', 'DATETIME'),
            ]
            with db.engine.connect() as conn:
                for col_name, col_type in new_g_cols:
                    if col_name not in g_cols:
                        try:
                            conn.execute(text(f'ALTER TABLE guarantors ADD COLUMN {col_name} {col_type}'))
                            conn.commit()
                        except Exception:
                            pass
        except Exception as e:
            print(f'Migration warning: {e}')
        try:
            from dateutil.relativedelta import relativedelta as _r
            seed_demo_data()
        except Exception as e:
            print(f'Seed error: {e}')

    return app


# For direct run / import
app = create_app()
